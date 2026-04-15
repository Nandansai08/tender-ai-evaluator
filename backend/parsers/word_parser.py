"""Word document parser using python-docx."""


def parse_word(file_path: str) -> dict:
    """
    Parse a .docx Word document and extract text content.

    Returns:
        dict with keys: text, paragraphs, tables, metadata, error
    """
    try:
        from docx import Document
    except ImportError:
        return {
            "text": "",
            "paragraphs": [],
            "tables": [],
            "metadata": {},
            "error": "python-docx is not installed. Install with: pip install python-docx",
        }

    try:
        doc = Document(file_path)
    except Exception as exc:
        return {
            "text": "",
            "paragraphs": [],
            "tables": [],
            "metadata": {},
            "error": str(exc),
        }

    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    tables = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        tables.append(rows)

    # Include table text in the full text dump
    table_texts = []
    for table_rows in tables:
        for row in table_rows:
            table_texts.append(" | ".join(row))

    all_text_parts = paragraphs + table_texts
    full_text = "\n".join(all_text_parts)

    metadata = {}
    try:
        core = doc.core_properties
        metadata = {
            "title": core.title or "",
            "author": core.author or "",
            "created": str(core.created) if core.created else "",
        }
    except Exception:
        pass

    return {
        "text": full_text,
        "paragraphs": paragraphs,
        "tables": tables,
        "metadata": metadata,
        "error": None,
    }
